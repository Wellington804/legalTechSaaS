import io
import unittest

from docx import Document

from app.services.document_text import citation_chunks, extract_upload_text, mark_pdf_pages


class DocumentTextTests(unittest.TestCase):
    def test_extracts_utf8_and_docx_with_a_hard_bound(self):
        self.assertEqual(extract_upload_text("text/plain", "Texto jurídico".encode()), "Texto jurídico")
        document = Document()
        document.add_paragraph("Fato informado")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Documento citado"
        output = io.BytesIO()
        document.save(output)
        self.assertEqual(extract_upload_text(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document", output.getvalue()
        ), "Fato informado\nDocumento citado")

    def test_citations_are_exact_and_rank_relevant_passages(self):
        text = "Primeiro trecho sem relação.\n\nA obrigação venceu em 10 de agosto e não foi paga."
        chunks = citation_chunks(text, "Quando venceu a obrigação?")
        self.assertEqual(chunks[0]["label"], "D-N2")
        self.assertEqual(chunks[0]["locator"], "§ 2")
        self.assertIn("obrigação venceu", chunks[0]["excerpt"])
        self.assertIn(chunks[0]["excerpt"], text)

    def test_preserves_pdf_page_and_paragraph_locators(self):
        text = mark_pdf_pages("Primeira página.\fFato principal.\n\nComprovante anexo.")
        chunks = citation_chunks(text, "Onde está o comprovante?", source_prefix="DOC2")
        self.assertEqual(chunks[0]["label"], "DOC2-P2-N2")
        self.assertEqual(chunks[0]["page"], 2)
        self.assertEqual(chunks[0]["locator"], "p. 2, § 2")
        self.assertEqual(chunks[0]["excerpt"], "Comprovante anexo.")


if __name__ == "__main__":
    unittest.main()
