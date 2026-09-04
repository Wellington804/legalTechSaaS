"""Run: python -m unittest tests.test_brand_documents -v (PDF gate uses LibreOffice)."""
import io
import shutil
import unittest
import zipfile
from unittest.mock import patch

from docx import Document
from docx.shared import Mm, Pt
from PIL import Image, PngImagePlugin
from pypdf import PdfReader, PdfWriter

from app.services import brand_documents as brand


def image_bytes():
    image = Image.new("RGBA", (240, 100), (12, 30, 60, 255))
    output = io.BytesIO()
    metadata = PngImagePlugin.PngInfo()
    metadata.add_text("private-name", "client-name")
    image.save(output, format="PNG", pnginfo=metadata)
    return output.getvalue()


def docx_bytes():
    document = Document()
    document.styles["Normal"].font.name = "Liberation Serif"
    document.styles["Normal"].font.size = Pt(12)
    document.sections[0].top_margin = Mm(25)
    document.sections[0].header.paragraphs[0].text = "Advogado — OAB informada"
    document.sections[0].footer.paragraphs[0].text = "Contato confirmado"
    document.add_paragraph("Ação de exemplo, sem dados reais.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def change_zip(content, name, replacement):
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(content)) as source, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as result:
        for entry in source.infolist():
            if entry.filename != name:
                result.writestr(entry, source.read(entry.filename))
        result.writestr(name, replacement)
    return output.getvalue()


def pdf_bytes(customize=None, pages=1):
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=595, height=842)
    if customize:
        customize(writer)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


class BrandReferenceTests(unittest.TestCase):
    def test_composed_layers_reserve_body_area_and_ignore_watermark(self):
        settings = {"layout_mode": "composed", "paper_size": "A4", "margin_top_mm": 30, "margin_bottom_mm": 25,
                    "layout_layers": [
                        {"y_percent": 1, "height_percent": 14, "role": "decoration", "visible": True},
                        {"y_percent": 91, "height_percent": 4, "role": "contact", "visible": True},
                        {"y_percent": 30, "height_percent": 30, "role": "watermark", "visible": True},
                    ]}
        self.assertEqual(brand.required_content_margins(settings), (49, 31))

    def test_image_is_normalized_without_metadata_and_input_is_immutable(self):
        original = image_bytes()
        snapshot = bytes(original)
        mime, normalized, analysis = brand.validate_reference("logo.png", original, "logo")
        self.assertEqual(mime, "image/png")
        self.assertEqual(original, snapshot)
        self.assertNotIn(b"client-name", normalized)
        self.assertEqual(analysis["identified"]["width"], 240)
        self.assertEqual(analysis["estimated"]["dominant_colors"], ["#0C1E3C"])
        with Image.open(io.BytesIO(normalized)) as image:
            self.assertEqual(image.mode, "RGBA")

    def test_image_format_spoofing_and_nonimages_as_logos_are_rejected(self):
        for filename, content, kind in (("logo.jpg", image_bytes(), "logo"), ("logo.svg", b"<svg/>", "logo"), ("logo.pdf", pdf_bytes(), "logo"), ("logo.png", b"not an image", "reference")):
            with self.subTest(filename=filename), self.assertRaises(ValueError):
                brand.validate_reference(filename, content, kind)

    def test_docx_extracts_real_styles_without_modifying_reference(self):
        original = docx_bytes()
        mime, returned, analysis = brand.validate_reference("reference.docx", original, "reference")
        self.assertEqual(returned, original)
        self.assertIn("wordprocessingml", mime)
        self.assertIn("Liberation Serif", analysis["identified"]["fonts"])
        self.assertIn(12, analysis["identified"]["font_sizes_pt"])
        self.assertAlmostEqual(analysis["identified"]["margins_mm"]["top"], 25, places=1)
        self.assertIn("OAB", analysis["identified"]["header_text"][0])
        self.assertEqual(analysis["estimated"], {})

    def test_docx_active_content_external_links_and_traversal_are_rejected(self):
        original = docx_bytes()
        attacks = [
            ("word/vbaProject.bin", b"macro"),
            ("word/embeddings/object.bin", b"object"),
            ("../secret", b"outside"),
            ("word/_rels/document.xml.rels", b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="r1" Type="hyperlink" Target="https://example.com" TargetMode="External"/></Relationships>'),
            ("word/document.xml", b'<!DOCTYPE root [<!ENTITY example "secret">]><root>&example;</root>'),
            ("word/document.xml", '<!DOCTYPE root [<!ENTITY example "secret">]><root>&example;</root>'.encode("utf-16")),
            ("word/document.xml", b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:instrText>DDEAUTO dangerous</w:instrText></w:document>'),
        ]
        for name, replacement in attacks:
            with self.subTest(name=name, size=len(replacement)), self.assertRaises(ValueError):
                brand.validate_reference("reference.docx", change_zip(original, name, replacement), "reference")

    def test_docx_zip_bomb_and_invalid_zip_are_rejected(self):
        bomb = change_zip(docx_bytes(), "word/document.xml", b"a" * 2_000_000)
        for content in (bomb, b"PKnotazip"):
            with self.assertRaises(ValueError):
                brand.validate_reference("reference.docx", content, "reference")

    def test_pdf_analysis_preserves_original_and_bounds_pages(self):
        original = pdf_bytes(pages=3)
        mime, returned, analysis = brand.validate_reference("reference.pdf", original, "reference")
        self.assertEqual(mime, "application/pdf")
        self.assertEqual(original, returned)
        self.assertEqual(analysis["identified"]["pages"], 3)
        with self.assertRaises(ValueError):
            brand.validate_reference("reference.pdf", pdf_bytes(pages=201), "reference")

    @unittest.skipUnless(shutil.which("pdftoppm"), "Poppler required for selected PDF page rendering")
    def test_only_selected_pdf_page_is_rendered_for_visual_analysis(self):
        rendered = brand.render_reference_page(pdf_bytes(pages=2), "application/pdf", 2)
        with Image.open(io.BytesIO(rendered)) as image:
            self.assertEqual(image.format, "PNG")
            self.assertGreater(image.width, 1_000)
        with self.assertRaises(ValueError):
            brand.render_reference_page(pdf_bytes(pages=2), "application/pdf", 3)

    def test_pdf_actions_attachments_and_encryption_are_rejected(self):
        changes = [lambda writer: writer.add_js("app.alert('unsafe')"), lambda writer: writer.add_attachment("payload.txt", b"secret"), lambda writer: writer.encrypt("password")]
        for change in changes:
            with self.assertRaises(ValueError):
                brand.validate_reference("reference.pdf", pdf_bytes(change), "reference")

    def test_image_page_crop_is_bounded_and_sanitized(self):
        cropped = brand.crop_reference(image_bytes(), "image/png", 1, (10, 10, 50, 50))
        with Image.open(io.BytesIO(cropped)) as image:
            self.assertEqual(image.size, (120, 50))
        for crop in ((90, 0, 20, 20), (0, 0, 0, 20)):
            with self.assertRaises(ValueError):
                brand.crop_reference(image_bytes(), "image/png", 1, crop)


class BrandDocumentTests(unittest.TestCase):
    def test_faint_watermark_isolation_drops_dark_body_text(self):
        source = Image.new("RGB", (80, 40), "white")
        source.putpixel((20, 20), (220, 220, 220))
        source.putpixel((60, 20), (10, 10, 10))
        buffer = io.BytesIO(); source.save(buffer, format="PNG")
        with Image.open(io.BytesIO(brand.isolate_layer_image(buffer.getvalue(), faint_only=True))) as isolated:
            self.assertGreater(isolated.getpixel((20, 20))[3], 0)
            self.assertEqual(isolated.getpixel((60, 20))[3], 0)

    def test_composed_canvas_renders_editable_primitive_layers(self):
        settings = {"layout_mode": "composed", "paper_color": "#FFFFFF", "layout_layers": [
            {"id": "navy", "kind": "rectangle", "x_percent": 0, "y_percent": 0, "width_percent": 100, "height_percent": 10, "color": "#102A56", "opacity": 1, "z_index": 0, "page_scope": "all"},
            {"id": "gold", "kind": "line", "x_percent": 0, "y_percent": 12, "width_percent": 100, "height_percent": 1, "color": "#C7A24B", "opacity": 1, "z_index": 1, "page_scope": "all", "line_thickness_pt": 2},
        ]}
        with Image.open(io.BytesIO(brand.render_brand_canvas(settings, {}))) as image:
            self.assertEqual(image.getpixel((image.width // 2, 2))[:3], (16, 42, 86))
            self.assertEqual(image.getpixel((image.width // 2, image.height // 2))[:3], (255, 255, 255))

    def test_reconstructed_canvas_renders_paper_and_positioned_branding(self):
        settings = {"layout_mode": "reconstructed", "paper_color": "#EDE7DD", "header_text": "DOUGLAS SIMÕES", "footer_text": "RUA · BAIRRO · MACEIÓ", "watermark_text": "DS", "watermark_position": "center", "watermark_x_percent": 50, "watermark_y_percent": 52, "watermark_opacity": .1}
        output = brand.render_brand_canvas(settings, {})
        with Image.open(io.BytesIO(output)) as image:
            self.assertEqual(image.getpixel((0, 0))[:3], (237, 231, 221))
            self.assertNotEqual(image.getbbox(), None)

    def test_exact_background_hides_automatic_title_but_keeps_body_editable(self):
        background = Image.new("RGB", (595, 842), "#EDE7DD")
        buffer = io.BytesIO(); background.save(buffer, format="PNG")
        settings = {"layout_mode": "exact", "background_asset_id": "paper", "show_document_title": False, "page_numbers": False}
        output = brand._render_docx("Título automático", "Corpo editável", settings, {"paper": buffer.getvalue()}, "plain")
        document = Document(io.BytesIO(output))
        self.assertEqual([paragraph.text for paragraph in document.paragraphs], ["Corpo editável"])
        self.assertIn('behindDoc="1"', document.sections[0].header._element.xml)
    def test_plain_content_is_not_interpreted_or_rewritten(self):
        content = "# Texto literal\nAção — R$ 1.000,00\n\n<script>alert('texto')</script>"
        output = brand._render_docx("Documento", content, {}, {}, "plain")
        document = Document(io.BytesIO(output))
        self.assertEqual([paragraph.text for paragraph in document.paragraphs[1:]], content.split("\n"))
        self.assertEqual(document.styles["Normal"].font.name, "Liberation Serif")
        self.assertAlmostEqual(document.sections[0].page_width.mm, 210, places=1)
        self.assertEqual(document.core_properties.author, "")

    def test_markdown_heading_list_and_table_are_native_editable_elements(self):
        output = brand._render_docx("Documento", "# Pedidos\n1. **Primeiro**\n- Segundo\n\n| Item | Valor |\n| --- | --- |\n| Honorários | R$ 100 |", {"paper_size": "LETTER"}, {}, "markdown")
        document = Document(io.BytesIO(output))
        self.assertEqual(document.paragraphs[1].style.name, "Heading 1")
        self.assertEqual(document.paragraphs[2].text, "1. Primeiro")
        self.assertTrue(document.paragraphs[2].runs[-1].bold)
        self.assertEqual(document.tables[0].rows[1].cells[0].text, "Honorários")
        self.assertAlmostEqual(document.sections[0].page_width.mm, 215.9, places=1)

    def test_first_header_logo_watermark_and_page_fields_are_present(self):
        settings = {"header_text": "Escritório", "first_header_text": "Primeira página", "footer_text": "Contato", "different_first_page": True, "logo_asset_id": "logo", "watermark_asset_id": "watermark", "watermark_position": "diagonal", "watermark_opacity": 0.1}
        output = brand._render_docx("Documento", "Conteúdo", settings, {"logo": image_bytes(), "watermark": image_bytes()}, "plain")
        document = Document(io.BytesIO(output))
        self.assertTrue(document.sections[0].different_first_page_header_footer)
        self.assertIn("Primeira página", document.sections[0].first_page_header.paragraphs[0].text)
        self.assertIn('behindDoc="1"', document.sections[0].header._element.xml)
        self.assertIn("NUMPAGES", document.sections[0].footer._element.xml)
        self.assertIn("PAGE", document.sections[0].first_page_footer._element.xml)

    def test_renderer_rejects_missing_assets_invalid_fonts_and_invalid_content(self):
        for settings, content in (({"font_family": "https://remote/font.ttf"}, "Texto"), ({"logo_asset_id": "absent"}, "Texto"), ({}, "bad\x00text"), ({}, "x" * (brand.MAX_CONTENT_CHARS + 1)), ({"header_text": "x" * 1000, "margin_top_mm": 10}, "Texto")):
            with self.subTest(settings=settings), self.assertRaises(ValueError):
                brand._render_docx("Documento", content, settings, {}, "plain")

    def test_pdf_unavailable_fails_honestly_and_slot_is_released(self):
        with patch.object(brand, "_office_binary", return_value=None):
            self.assertFalse(brand.pdf_available())
            for _ in range(2):
                with self.assertRaisesRegex(RuntimeError, "indisponível"):
                    brand.render_documents("Documento", "Conteúdo", {}, {})

    @unittest.skipUnless(brand.pdf_available(), "LibreOffice + approved fonts required for actual PDF validation")
    def test_actual_multipage_pdf_has_selectable_text_first_header_and_watermark(self):
        settings = {"header_text": "Cabeçalho recorrente", "first_header_text": "Cabeçalho exclusivo inicial", "footer_text": "Contato profissional", "different_first_page": True, "watermark_text": "IDENTIDADE APROVADA", "watermark_opacity": 0.1, "watermark_position": "diagonal", "logo_asset_id": "logo"}
        content = "\n".join(f"Parágrafo {index}: ação, diligência e revisão humana. Texto preservado." for index in range(180))
        docx, pdf = brand.render_documents("Documento de validação", content, settings, {"logo": image_bytes()})
        self.assertTrue(docx.startswith(b"PK"))
        self.assertTrue(pdf.startswith(b"%PDF"))
        reader = PdfReader(io.BytesIO(pdf))
        self.assertGreater(len(reader.pages), 3)
        first = reader.pages[0].extract_text()
        last = reader.pages[-1].extract_text()
        self.assertIn("Cabeçalho exclusivo inicial", first)
        self.assertNotIn("Cabeçalho exclusivo inicial", last)
        self.assertIn("Cabeçalho recorrente", last)
        self.assertIn("Parágrafo 179", last)
        self.assertIn("Contato profissional", last)
        self.assertIn("Página", last)
        self.assertGreater(len(reader.pages[0].images), 1)


if __name__ == "__main__":
    unittest.main()
