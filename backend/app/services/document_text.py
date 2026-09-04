"""Bounded text extraction for already validated office uploads."""
import io
import re
import unicodedata
import zipfile

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from openpyxl import load_workbook


MAX_PAGES = 300
MAX_TEXT_CHARS = 250_000
MAX_DOCX_ENTRIES = 2_000
MAX_DOCX_UNCOMPRESSED = 40 * 1024 * 1024
PAGE_MARKER = "[[LEXFLOW_PAGE:{page}]]"
PAGE_MARKER_RE = re.compile(r"^\[\[LEXFLOW_PAGE:(\d+)\]\]$", re.MULTILINE)


class TextExtractionError(ValueError):
    pass


def mark_pdf_pages(text: str) -> str:
    """Preserve OCR page boundaries without needing a parallel source-map table."""
    pages = text.replace("\r\n", "\n").split("\f")
    return "\n\n".join(
        f"{PAGE_MARKER.format(page=index)}\n{page.strip()}"
        for index, page in enumerate(pages, 1)
        if page.strip()
    )


def _bounded(texts) -> str | None:
    result: list[str] = []
    size = 0
    for value in texts:
        text = str(value or "").replace("\x00", "").strip()
        if not text:
            continue
        remaining = MAX_TEXT_CHARS - size
        if remaining <= 0:
            break
        result.append(text[:remaining])
        size += len(result[-1]) + 1
    joined = "\n".join(result).strip()
    return joined or None


def extract_upload_text(content_type: str, content: bytes) -> str | None:
    try:
        if content_type == "text/plain":
            return _bounded([content.decode("utf-8")])
        if content_type == "application/pdf":
            reader = PdfReader(io.BytesIO(content), strict=True)
            if reader.is_encrypted or len(reader.pages) > MAX_PAGES:
                raise TextExtractionError("PDF protegido ou extenso demais para indexação.")
            return _bounded(
                f"{PAGE_MARKER.format(page=index)}\n{page.extract_text() or ''}"
                for index, page in enumerate(reader.pages, 1)
            )
        if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > MAX_DOCX_ENTRIES or sum(item.file_size for item in entries) > MAX_DOCX_UNCOMPRESSED:
                    raise TextExtractionError("DOCX extenso demais para indexação.")
                if any(item.compress_size and item.file_size / item.compress_size > 200 for item in entries):
                    raise TextExtractionError("Compactação do DOCX não é segura para indexação.")
            document = Document(io.BytesIO(content))
            return _bounded([*(paragraph.text for paragraph in document.paragraphs),
                             *(cell.text for table in document.tables for row in table.rows for cell in row.cells)])
        if content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            try:
                if len(workbook.sheetnames) > 100:
                    raise TextExtractionError("Planilha extensa demais para indexacao.")
                return _bounded(
                    cell.value
                    for sheet in workbook.worksheets
                    for row_number, row in enumerate(sheet.iter_rows(), 1)
                    if row_number <= 20_000
                    for cell in row
                )
            finally:
                workbook.close()
    except TextExtractionError:
        raise
    except (OSError, ValueError, KeyError, TypeError, zipfile.BadZipFile, PdfReadError) as exc:
        raise TextExtractionError("Não foi possível extrair texto seguro do arquivo.") from exc
    return None


def citation_chunks(text: str, question: str, limit: int = 6, *, source_prefix: str = "D") -> list[dict]:
    """Return exact, bounded excerpts ranked by query terms; never invent a source."""
    normalized = unicodedata.normalize("NFKD", question).encode("ascii", "ignore").decode().lower()
    terms = {term for term in re.findall(r"[a-z0-9]{4,}", normalized)}
    source_prefix = re.sub(r"[^A-Z0-9_-]", "", source_prefix.upper())[:24] or "D"
    paragraphs: list[tuple[int | None, int, str]] = []
    matches = list(PAGE_MARKER_RE.finditer(text))
    if matches:
        for position, match in enumerate(matches):
            end = matches[position + 1].start() if position + 1 < len(matches) else len(text)
            page = int(match.group(1))
            for paragraph_index, part in enumerate(re.split(r"\n{2,}", text[match.end():end]), 1):
                if part.strip():
                    paragraphs.append((page, paragraph_index, part.strip()))
    else:
        paragraphs = [(None, index, part.strip()) for index, part in enumerate(re.split(r"\n{2,}", text), 1) if part.strip()]
    chunks: list[tuple[int | None, int, int, str]] = []
    for page, paragraph_index, paragraph in paragraphs:
        for start in range(0, len(paragraph), 1_200):
            chunks.append((page, paragraph_index, start // 1_200 + 1, paragraph[start:start + 1_200]))
    ranked = sorted(enumerate(chunks), key=lambda item: (
        -sum(term in unicodedata.normalize("NFKD", item[1][3]).encode("ascii", "ignore").decode().lower() for term in terms),
        item[0],
    ))[:max(1, min(limit, 40))]
    result = []
    for _index, (page, paragraph, chunk, excerpt) in ranked:
        suffix = f"P{page}-N{paragraph}" if page is not None else f"N{paragraph}"
        if chunk > 1:
            suffix += f"-C{chunk}"
        result.append({
            "label": f"{source_prefix}-{suffix}",
            "page": page,
            "paragraph": paragraph,
            "locator": f"p. {page}, § {paragraph}" if page is not None else f"§ {paragraph}",
            "excerpt": excerpt,
        })
    return result
