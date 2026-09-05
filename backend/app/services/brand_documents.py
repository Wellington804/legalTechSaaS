"""Private reference inspection and native DOCX/PDF rendering.

Uploaded Office/PDF documents are inspected, never opened by LibreOffice. Only a
new document built from validated text, settings and normalized images is rendered.
"""
from __future__ import annotations

import io
import math
import os
import re
import shutil
import signal
import subprocess
import tempfile
import threading
import warnings
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image, ImageChops, ImageDraw, ImageEnhance, ImageFont, ImageOps
from pypdf import PdfReader, filters
from pypdf.generic import ArrayObject, DictionaryObject, IndirectObject
from app.schemas.branding import FONT_FAMILIES

FONTS = FONT_FAMILIES
_FONT_FILES = {
    "Liberation Serif": "LiberationSerif-Regular.ttf",
    "Liberation Sans": "LiberationSans-Regular.ttf",
    "Liberation Mono": "LiberationMono-Regular.ttf",
    "DejaVu Serif": "DejaVuSerif.ttf",
    "DejaVu Sans": "DejaVuSans.ttf",
    "DejaVu Sans Mono": "DejaVuSansMono.ttf",
    "Noto Serif": "NotoSerif-Regular.ttf",
    "Noto Sans": "NotoSans-Regular.ttf",
    "Noto Mono": "NotoMono-Regular.ttf",
    "Carlito": "Carlito-Regular.ttf",
    "Caladea": "Caladea-Regular.ttf",
    "Lato": "Lato-Regular.ttf",
    "Tinos": "Tinos-Regular.ttf",
    "Arial": "Arimo-Regular.ttf",
    "Calibri": "Carlito-Regular.ttf",
    "Cambria": "Caladea-Regular.ttf",
    "Courier New": "Cousine-Regular.ttf",
    "Times New Roman": "Tinos-Regular.ttf",
    "DejaVu Sans Condensed": "DejaVuSansCondensed.ttf",
    "DejaVu Serif Condensed": "DejaVuSerifCondensed.ttf",
    "Noto Sans Display": "NotoSansDisplay-Regular.ttf",
    "Noto Serif Display": "NotoSerifDisplay-Regular.ttf",
    "Noto Sans Mono": "NotoSansMono-Regular.ttf",
}
_FONT_DIRECTORIES = (
    Path("/usr/share/fonts/truetype/liberation2"), Path("/usr/share/fonts/truetype/liberation"),
    Path("/usr/share/fonts/truetype/dejavu"), Path("/usr/share/fonts/truetype/noto"),
    Path("/usr/share/fonts/truetype/lato"), Path("/usr/share/fonts/truetype/crosextra"),
    Path("/usr/share/fonts/truetype/croscore"),
    Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts",
)
MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_OUTPUT_BYTES = 25 * 1024 * 1024
MAX_CONTENT_CHARS = 300_000
MAX_PAGES = 200
MAX_XML_BYTES = 5 * 1024 * 1024
MAX_UNPACKED_BYTES = 40 * 1024 * 1024
MAX_IMAGE_PIXELS = 16_000_000
RENDER_TIMEOUT = 45
REFERENCE_DPI = 144
_RENDER_SLOTS = threading.BoundedSemaphore(1)
_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}

# pypdf's default decompression budget is 75 MB per stream, excessive for a
# branding sample. No external image decoder is allowed during inspection.
for _limit in ("ZLIB_MAX_OUTPUT_LENGTH", "LZW_MAX_OUTPUT_LENGTH", "RUN_LENGTH_MAX_OUTPUT_LENGTH", "MAX_DECLARED_STREAM_LENGTH", "MAX_ARRAY_BASED_STREAM_OUTPUT_LENGTH", "IMAGE_MAX_BUFFER_SIZE"):
    setattr(filters, _limit, MAX_UPLOAD_BYTES)
filters.JBIG2DEC_BINARY = None


def _xml(content: bytes) -> ET.Element:
    if len(content) > MAX_XML_BYTES or re.search(br"<!\s*(?:DOCTYPE|ENTITY)", content, re.I):
        raise ValueError("XML com entidades ou tamanho não permitido.")
    try:
        class NoDtd(ET.TreeBuilder):
            def doctype(self, name, pubid, system):
                raise ValueError("DTD não permitido na referência.")

        root = ET.fromstring(content, parser=ET.XMLParser(target=NoDtd()))
    except ET.ParseError as exc:
        raise ValueError("XML inválido na referência.") from exc
    if sum(1 for _ in root.iter()) > 100_000:
        raise ValueError("Referência XML excessivamente complexa.")
    return root


def _inspect_docx(content: bytes) -> dict:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            names = [entry.filename for entry in entries]
            if len(entries) > 1000 or len(names) != len(set(names)):
                raise ValueError("Pacote Word inválido ou excessivamente complexo.")
            if not {"[Content_Types].xml", "word/document.xml"}.issubset(names):
                raise ValueError("O arquivo não é um documento DOCX.")
            if sum(entry.file_size for entry in entries) > MAX_UNPACKED_BYTES:
                raise ValueError("Conteúdo descompactado excede o limite.")
            xml_parts = {}
            for entry in entries:
                path = PurePosixPath(entry.filename)
                lower = entry.filename.lower()
                if path.is_absolute() or ".." in path.parts or "\\" in lower or ":" in lower or entry.flag_bits & 1:
                    raise ValueError("Pacote Word contém caminho ou proteção não permitida.")
                if entry.file_size > MAX_UPLOAD_BYTES or entry.file_size > max(entry.compress_size, 1) * 200:
                    raise ValueError("Compressão excessiva na referência.")
                if any(token in lower for token in ("vbaproject", "activex", "embeddings/", "oleobject", "customui/")):
                    raise ValueError("Macros e objetos incorporados não são permitidos.")
                if lower.endswith((".xml", ".rels")):
                    root = _xml(archive.read(entry))
                    xml_parts[entry.filename] = root
                    for element in root.iter():
                        local = element.tag.rsplit("}", 1)[-1]
                        if local in {"altChunk", "object", "oleObject", "control", "attachedTemplate"}:
                            raise ValueError("Conteúdo ativo não permitido no documento.")
                        if local == "Relationship" and (element.get("TargetMode", "").lower() == "external" or re.match(r"(?i)^[a-z][a-z0-9+.-]*:", element.get("Target", ""))):
                            raise ValueError("Referências externas não são permitidas.")
                        if local in {"instrText", "fldSimple"}:
                            instruction = (element.text or "") + element.get(_W + "instr", "")
                            if re.search(r"\b(?:DDE|DDEAUTO|INCLUDETEXT|INCLUDEPICTURE|LINK|DATABASE)\b", instruction, re.I):
                                raise ValueError("Campos externos não são permitidos.")
                        if local in {"Default", "Override"} and any(token in element.get("ContentType", "").lower() for token in ("macroenabled", "vba", "activex", "oleobject")):
                            raise ValueError("Tipo ativo não permitido no documento.")
            fonts, sizes, colors = set(), set(), set()
            for root in xml_parts.values():
                for item in root.iter():
                    if item.tag == _W + "rFonts":
                        fonts.update(value[:100] for key, value in item.attrib.items() if key in {_W + "ascii", _W + "hAnsi", _W + "cs", _W + "eastAsia"})
                    if item.tag == _W + "sz" and item.get(_W + "val", "").isdigit():
                        sizes.add(int(item.get(_W + "val")) / 2)
                    if item.tag == _W + "color" and re.fullmatch(r"[0-9A-Fa-f]{6}", item.get(_W + "val", "")):
                        colors.add("#" + item.get(_W + "val").upper())
            document = xml_parts["word/document.xml"]
            margins = document.find(".//" + _W + "pgMar")
            identified = {"format": "DOCX", "fonts": sorted(fonts)[:40], "font_sizes_pt": sorted(sizes)[:40], "colors": sorted(colors)[:40]}
            if margins is not None:
                identified["margins_mm"] = {key: round(int(margins.get(_W + key)) * 25.4 / 1440, 2) for key in ("top", "bottom", "left", "right") if (margins.get(_W + key) or "").isdigit()}
            for part, key in (("header", "header_text"), ("footer", "footer_text")):
                texts = [" ".join(node.text or "" for node in root.iter(_W + "t"))[:1000] for name, root in sorted(xml_parts.items()) if re.fullmatch(rf"word/{part}\d+\.xml", name)]
                if texts:
                    identified[key] = texts
            return {"identified": identified, "estimated": {}, "warnings": ["Estilos declarados não garantem a mesma fonte instalada. Confirme os elementos antes de publicar."]}
    except (zipfile.BadZipFile, RuntimeError, OSError) as exc:
        raise ValueError("Pacote DOCX inválido.") from exc


def _inspect_pdf(content: bytes) -> dict:
    try:
        reader = PdfReader(io.BytesIO(content), strict=True, root_object_recovery_limit=1000)
        if reader.is_encrypted:
            raise ValueError("PDF protegido por senha não é permitido.")
        if sum(len(items) for items in reader.xref.values()) > 20_000 or len(reader.xref_objStm) > 20_000:
            raise ValueError("PDF excessivamente complexo.")
        stack, seen, visited = [reader.trailer], set(), 0
        forbidden = {"/JavaScript", "/JS", "/AA", "/OpenAction", "/A", "/Launch", "/EmbeddedFiles", "/EF", "/RichMedia", "/XFA", "/AcroForm", "/Collection"}
        while stack:
            obj = stack.pop()
            visited += 1
            if visited > 100_000:
                raise ValueError("PDF excessivamente complexo.")
            if isinstance(obj, IndirectObject):
                identity = (obj.idnum, obj.generation)
                if identity in seen:
                    continue
                seen.add(identity)
                obj = obj.get_object()
            if isinstance(obj, DictionaryObject):
                if forbidden.intersection(obj) or obj.get("/Type") in {"/Filespec", "/EmbeddedFile", "/Action"} or obj.get("/S") in {"/URI", "/GoToR", "/GoToE", "/Launch", "/JavaScript", "/SubmitForm", "/ImportData", "/Rendition", "/Movie", "/Sound"}:
                    raise ValueError("Ações, formulários e anexos incorporados não são permitidos no PDF.")
                stack.extend(obj.values())
            elif isinstance(obj, ArrayObject):
                stack.extend(obj)
        if not 1 <= len(reader.pages) <= MAX_PAGES:
            raise ValueError(f"PDF deve conter entre 1 e {MAX_PAGES} páginas.")
        fonts = set()
        page_sizes = set()
        for page in reader.pages:
            width_mm, height_mm = float(page.mediabox.width) * 25.4 / 72, float(page.mediabox.height) * 25.4 / 72
            if not all(math.isfinite(value) and 20 <= value <= 600 for value in (width_mm, height_mm)):
                raise ValueError("PDF possui página com dimensões não permitidas.")
            page_sizes.add((round(width_mm, 2), round(height_mm, 2)))
            resources = page.get("/Resources", {}).get_object() if page.get("/Resources") else {}
            font_resources = resources.get("/Font", {}).get_object() if resources.get("/Font") else {}
            for font in font_resources.values():
                name = str(font.get_object().get("/BaseFont", ""))[1:]
                if name:
                    fonts.add(re.sub(r"^[A-Z]{6}\+", "", name)[:100])
        return {"identified": {"format": "PDF", "pages": len(reader.pages), "fonts": sorted(fonts)[:40], "page_sizes_mm": sorted(page_sizes)}, "estimated": {}, "warnings": ["A referência não foi convertida. Margens, cores e hierarquia visual exigem revisão; fontes de imagens não podem ser identificadas com exatidão."]}
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("PDF inválido ou fora dos limites de processamento.") from exc


def _normalize_image(content: bytes, expected: str) -> tuple[bytes, dict]:
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(io.BytesIO(content), formats=["PNG", "JPEG"]) as source:
                if source.format != expected or source.width * source.height > MAX_IMAGE_PIXELS or max(source.size) > 8000 or getattr(source, "n_frames", 1) != 1:
                    raise ValueError("Imagem incompatível, animada ou acima do limite de resolução.")
                source.load()
                normalized = ImageOps.exif_transpose(source).convert("RGBA" if expected == "PNG" else "RGB")
                normalized.info.clear()
                buffer = io.BytesIO()
                normalized.save(buffer, format=expected, **({"quality": 95} if expected == "JPEG" else {}))
                if buffer.tell() > MAX_UPLOAD_BYTES:
                    raise ValueError("Imagem normalizada excede o limite de tamanho.")
                sample = normalized.convert("RGB")
                sample.thumbnail((128, 128))
                palette = sample.quantize(colors=5).convert("RGB").getcolors(16384) or []
                dominant = ["#%02X%02X%02X" % color for _, color in sorted(palette, reverse=True)[:5]]
                return buffer.getvalue(), {"identified": {"format": expected, "width": normalized.width, "height": normalized.height}, "estimated": {"dominant_colors": dominant}, "warnings": ["Cores dominantes são estimativas. Tipografia de imagem precisa de confirmação."]}
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("Imagem inválida ou fora dos limites de processamento.") from exc


def validate_reference(filename: str, content: bytes, kind: str) -> tuple[str, bytes, dict]:
    """Inspect original documents or return a metadata-free normalized image."""
    if kind not in {"reference", "logo", "logo_dark", "logo_mono", "watermark", "background"}:
        raise ValueError("Tipo de referência inválido.")
    if not content or len(content) > MAX_UPLOAD_BYTES:
        raise ValueError("Referência vazia ou maior que 10 MB.")
    extension = Path(filename).suffix.lower()
    if extension in {".png", ".jpg", ".jpeg"}:
        expected = "PNG" if extension == ".png" else "JPEG"
        normalized, analysis = _normalize_image(content, expected)
        return ("image/png" if expected == "PNG" else "image/jpeg"), normalized, analysis
    if kind != "reference":
        raise ValueError("Elementos visuais devem ser PNG ou JPEG.")
    if extension == ".docx":
        return _DOCX_MIME, content, _inspect_docx(content)
    if extension == ".pdf" and content.startswith(b"%PDF-"):
        return "application/pdf", content, _inspect_pdf(content)
    raise ValueError("Formatos aceitos: PNG, JPEG, DOCX e PDF sem conteúdo ativo.")


def render_reference_page(content: bytes, content_type: str, page: int = 1) -> bytes:
    """Return one sanitized PNG page; the complete PDF is never sent to visual AI."""
    if isinstance(page, bool) or not isinstance(page, int) or page < 1:
        raise ValueError("Página de referência inválida.")
    if content_type in {"image/png", "image/jpeg"}:
        if page != 1:
            raise ValueError("A imagem possui somente uma página.")
        extension = ".png" if content_type == "image/png" else ".jpg"
        _, safe, _ = validate_reference("referencia" + extension, content, "reference")
        with Image.open(io.BytesIO(safe)) as source:
            result = io.BytesIO()
            source.convert("RGB").save(result, format="PNG", optimize=True)
            return result.getvalue()
    if content_type != "application/pdf":
        raise ValueError("Somente páginas de PDF ou imagem podem ser visualizadas.")
    analysis = _inspect_pdf(content)
    if page > analysis["identified"]["pages"]:
        raise ValueError("A página escolhida não existe nesta referência.")
    binary = shutil.which("pdftoppm")
    if not binary:
        raise RuntimeError("O servidor não possui o renderizador seguro de páginas PDF.")
    with tempfile.TemporaryDirectory(prefix="lexflow-reference-") as directory:
        work = Path(directory)
        source, target = work / "reference.pdf", work / "page"
        source.write_bytes(content)
        command = [binary, "-f", str(page), "-l", str(page), "-singlefile", "-r", str(REFERENCE_DPI), "-png", str(source), str(target)]
        if os.name == "posix":
            limiter = shutil.which("prlimit")
            if not limiter:
                raise RuntimeError("O servidor precisa de prlimit para renderizar referências.")
            command = [limiter, "--cpu=20", "--as=536870912", f"--fsize={MAX_UPLOAD_BYTES}", "--nofile=128", "--"] + command
        environment = {key: value for key, value in os.environ.items() if key in {"PATH", "SYSTEMROOT", "WINDIR", "LANG", "LC_ALL"}}
        try:
            process = subprocess.run(command, cwd=directory, env=environment, stdin=subprocess.DEVNULL,
                                     stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=25, check=False,
                                     creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0)
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise RuntimeError("Não foi possível renderizar a página de referência com segurança.") from exc
        output = target.with_suffix(".png")
        if process.returncode or not output.is_file() or not 0 < output.stat().st_size <= MAX_UPLOAD_BYTES:
            raise RuntimeError("A página de referência não passou na renderização segura.")
        normalized, _ = _normalize_image(output.read_bytes(), "PNG")
        return normalized


def crop_reference(content: bytes, content_type: str, page: int, crop: tuple[float, float, float, float]) -> bytes:
    rendered = render_reference_page(content, content_type, page)
    x, y, width, height = crop
    if any(isinstance(value, bool) or not isinstance(value, (int, float)) or not math.isfinite(value) for value in crop):
        raise ValueError("Área de recorte inválida.")
    if x < 0 or y < 0 or width <= 0 or height <= 0 or x + width > 100 or y + height > 100:
        raise ValueError("A área escolhida precisa ficar dentro da página.")
    with Image.open(io.BytesIO(rendered)) as source:
        left, top = round(source.width * x / 100), round(source.height * y / 100)
        right, bottom = round(source.width * (x + width) / 100), round(source.height * (y + height) / 100)
        if right - left < 20 or bottom - top < 20:
            raise ValueError("A área escolhida é pequena demais.")
        result = io.BytesIO()
        source.crop((left, top, right, bottom)).save(result, format="PNG", optimize=True)
        normalized, _ = _normalize_image(result.getvalue(), "PNG")
        return normalized


def isolate_layer_image(content: bytes, *, faint_only: bool = False) -> bytes:
    """Remove a near-uniform crop background without accepting executable/vector content."""
    with Image.open(io.BytesIO(content), formats=["PNG"]) as source:
        image = source.convert("RGBA")
        corners = [image.convert("RGB").getpixel(point) for point in ((0, 0), (image.width - 1, 0), (0, image.height - 1), (image.width - 1, image.height - 1))]
        background = tuple(sorted(pixel[channel] for pixel in corners)[len(corners) // 2] for channel in range(3))
        rgb = image.convert("RGB")
        difference = ImageChops.difference(rgb, Image.new("RGB", image.size, background))
        channels = difference.split()
        distance = ImageChops.lighter(ImageChops.lighter(channels[0], channels[1]), channels[2])
        alpha = distance.point(lambda value: 0 if value <= 4 or faint_only and value > 110 else min(255, (value - 4) * 18))
        if image.getchannel("A").getextrema() != (255, 255):
            alpha = ImageChops.multiply(alpha, image.getchannel("A"))
        image.putalpha(alpha)
        output = io.BytesIO()
        image.save(output, format="PNG", optimize=True)
        normalized, _ = _normalize_image(output.getvalue(), "PNG")
        return normalized


def _font_file(family: str) -> str | None:
    stem = _FONT_FILES.get(family)
    if not stem:
        return None
    for directory in _FONT_DIRECTORIES:
        candidate = directory / stem
        if candidate.is_file():
            return str(candidate)
    return None


def _office_binary() -> str | None:
    return shutil.which("libreoffice") or shutil.which("soffice")


def pdf_available() -> bool:
    return bool(_office_binary() and all(_font_file(font) for font in FONTS) and (os.name != "posix" or shutil.which("prlimit")))


def _number(settings: dict, key: str, default: float, low: float, high: float) -> float:
    value = settings.get(key, default)
    if isinstance(value, bool) or not isinstance(value, (int, float)) or not math.isfinite(value) or not low <= value <= high:
        raise ValueError(f"Configuração inválida: {key}.")
    return float(value)


def _color(settings: dict, key: str, default: str) -> RGBColor:
    value = settings.get(key, default)
    if not isinstance(value, str) or not re.fullmatch(r"#[0-9a-fA-F]{6}", value):
        raise ValueError(f"Cor inválida: {key}.")
    return RGBColor.from_string(value[1:])


def _text(value: str, limit: int) -> str:
    if not isinstance(value, str) or len(value) > limit or re.search(r"[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff]", value):
        raise ValueError("Texto contém caracteres inválidos ou excede o limite.")
    return value


def _field(paragraph, name: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), name)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _font(style, family: str) -> None:
    style.font.name = family
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for key in list(fonts.attrib):
        if key.endswith("Theme"):
            del fonts.attrib[key]
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn("w:" + key), family)


def _letter_spacing(style, points: float) -> None:
    spacing = style.element.get_or_add_rPr().find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        style.element.get_or_add_rPr().append(spacing)
    spacing.set(qn("w:val"), str(round(points * 20)))


def _paragraph_border(paragraph, edge: str, color: str, thickness_pt: float = .75) -> None:
    borders = paragraph._p.get_or_add_pPr().find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph._p.get_or_add_pPr().append(borders)
    border = OxmlElement(f"w:{edge}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(round(thickness_pt * 8)))
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color.lstrip("#"))
    borders.append(border)


def _inline(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*[^*\n]+\*\*)", text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def _rgb(value: str) -> tuple[int, int, int]:
    if not isinstance(value, str) or not re.fullmatch(r"#[0-9a-fA-F]{6}", value):
        raise ValueError("Cor inválida no timbrado.")
    return tuple(int(value[index:index + 2], 16) for index in (1, 3, 5))


def _tracked_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.FreeTypeFont,
                  fill: tuple[int, int, int, int], align: str, spacing_px: float = 0, uppercase: bool = False) -> None:
    text = _text(text, 1000).upper() if uppercase else _text(text, 1000)
    if not text:
        return
    left, top, right, _bottom = box
    line_height = max(font.getbbox("Ag")[3] - font.getbbox("Ag")[1], 1) * 1.35
    lines = []
    for raw_line in text.splitlines():
        current = ""
        for word in raw_line.split(" "):
            candidate = word if not current else current + " " + word
            width = draw.textlength(candidate, font=font) + max(0, len(candidate) - 1) * spacing_px
            if current and width > right - left:
                lines.append(current)
                current = word
            else:
                current = candidate
        lines.append(current)
    for line_number, line in enumerate(lines):
        widths = [draw.textlength(char, font=font) for char in line]
        line_width = sum(widths) + max(0, len(line) - 1) * spacing_px
        x = left if align == "left" else right - line_width if align == "right" else left + (right - left - line_width) / 2
        y = top + line_number * line_height
        for char, width in zip(line, widths):
            draw.text((round(x), round(y)), char, font=font, fill=fill)
            x += width + spacing_px


def _draw_layer_icon(draw: ImageDraw.ImageDraw, icon: str, box: tuple[int, int, int, int], fill, width: int) -> None:
    """Draw approved contact symbols with primitives so PDF, DOCX and browser semantics stay independent."""
    left, top, right, bottom = box
    cx, cy = (left + right) / 2, (top + bottom) / 2
    inset = max(1, round(min(right - left, bottom - top) * .14))
    if icon == "email":
        draw.rounded_rectangle((left + inset, top + inset * 2, right - inset, bottom - inset * 2), radius=inset, outline=fill, width=width)
        draw.line((left + inset, top + inset * 2, cx, cy, right - inset, top + inset * 2), fill=fill, width=width)
    elif icon == "location":
        radius = max(2, round((right - left) * .28))
        draw.ellipse((cx - radius, top + inset, cx + radius, top + inset + radius * 2), outline=fill, width=width)
        draw.ellipse((cx - width, top + inset + radius - width, cx + width, top + inset + radius + width), fill=fill)
        draw.polygon(((cx - radius, top + inset + radius * 1.5), (cx + radius, top + inset + radius * 1.5), (cx, bottom - inset)), outline=fill)
    elif icon == "website":
        draw.ellipse((left + inset, top + inset, right - inset, bottom - inset), outline=fill, width=width)
        draw.arc((cx - (right - left) * .2, top + inset, cx + (right - left) * .2, bottom - inset), 90, 270, fill=fill, width=width)
        draw.line((left + inset, cy, right - inset, cy), fill=fill, width=width)
    elif icon in {"phone", "whatsapp"}:
        if icon == "whatsapp":
            draw.ellipse((left + inset, top + inset, right - inset, bottom - inset), outline=fill, width=width)
        phone_box = (left + inset * 2, top + inset * 2, right - inset * 2, bottom - inset * 2)
        draw.arc(phone_box, 120, 245, fill=fill, width=max(width, 2))
        draw.arc(phone_box, 300, 65, fill=fill, width=max(width, 2))


def _render_composed(canvas: Image.Image, settings: dict, assets: dict[str, bytes], *, first_page: bool, dpi: int) -> None:
    size = canvas.size
    for layer in sorted(settings.get("layout_layers", []), key=lambda item: (item.get("z_index", 0), item.get("id", ""))):
        if not layer.get("visible", True):
            continue
        scope = layer.get("page_scope", "all")
        if scope == "first" and not first_page or scope == "continuation" and first_page:
            continue
        x = round(size[0] * float(layer["x_percent"]) / 100)
        y = round(size[1] * float(layer["y_percent"]) / 100)
        width = max(1, round(size[0] * float(layer["width_percent"]) / 100))
        height = max(1, round(size[1] * float(layer["height_percent"]) / 100))
        surface = Image.new("RGBA", (width, height))
        draw = ImageDraw.Draw(surface)
        fill = _rgb(layer.get("color", "#17324D")) + (255,)
        kind = layer.get("kind")
        if kind == "rectangle":
            draw.rectangle((0, 0, width, height), fill=fill)
        elif kind == "line":
            thickness = max(1, round(_number(layer, "line_thickness_pt", 1, .25, 12) * dpi / 72))
            draw.line((0, height / 2, width, height / 2), fill=fill, width=thickness)
        elif kind == "image":
            asset_id = str(layer.get("asset_id") or "")
            if not asset_id or asset_id not in assets:
                raise ValueError("Imagem de camada não encontrada.")
            raw = assets[asset_id]
            _, safe, _ = validate_reference("layer.png" if raw.startswith(b"\x89PNG") else "layer.jpg", raw, "logo")
            with Image.open(io.BytesIO(safe)) as source:
                picture = source.convert("RGBA")
                picture = ImageEnhance.Contrast(picture).enhance(_number(layer, "image_contrast", 1, .5, 3))
                picture.thumbnail((width, height), Image.Resampling.LANCZOS)
                surface.alpha_composite(picture, (round((width - picture.width) / 2), round((height - picture.height) / 2)))
        elif kind in {"text", "icon_text"}:
            font_path = _font_file(layer.get("font_family", settings.get("utility_font_family", FONTS[0])))
            if not font_path:
                raise RuntimeError("A fonte da camada não está instalada no servidor.")
            font_size = max(1, round(_number(layer, "font_size_pt", 8, 5, 40) * dpi / 72))
            font = ImageFont.truetype(font_path, font_size)
            icon = layer.get("icon", "none") if kind == "icon_text" else "none"
            text_left = 0
            if icon != "none":
                icon_size = min(height, max(font_size * 2, round(width * .12)))
                _draw_layer_icon(draw, icon, (0, 0, icon_size, min(height, icon_size)), fill, max(1, round(dpi / 96)))
                text_left = icon_size + max(2, round(font_size * .35))
            text = _text(layer.get("text", ""), 500)
            _tracked_text(draw, (text_left, 0, width, height), text,
                          font, fill, layer.get("alignment", "left"),
                          _number(layer, "letter_spacing_pt", 0, 0, 5) * dpi / 72,
                          bool(layer.get("uppercase", False)))
            if layer.get("font_weight") == "bold" and text:
                _tracked_text(draw, (text_left + 1, 0, width, height), text,
                              font, fill, layer.get("alignment", "left"),
                              _number(layer, "letter_spacing_pt", 0, 0, 5) * dpi / 72,
                              bool(layer.get("uppercase", False)))
        else:
            raise ValueError("Tipo de camada inválido.")
        opacity = _number(layer, "opacity", 1, 0, 1)
        if opacity < 1:
            surface.putalpha(surface.getchannel("A").point(lambda alpha: round(alpha * opacity)))
        rotation = _number(layer, "rotation_deg", 0, -180, 180)
        if rotation:
            surface = surface.rotate(-rotation, expand=True, resample=Image.Resampling.BICUBIC)
            x += round((width - surface.width) / 2)
            y += round((height - surface.height) / 2)
        canvas.alpha_composite(surface, (x, y))


def required_content_margins(settings: dict) -> tuple[float, float]:
    """Return safe top/bottom margins for the visible fixed design plane."""
    top = float(settings.get("margin_top_mm", 30))
    bottom = float(settings.get("margin_bottom_mm", 25))
    if settings.get("layout_mode") != "composed":
        return top, bottom
    height_mm = 297.0 if settings.get("paper_size", "A4") == "A4" else 279.4
    header_bottom = footer_top = None
    for layer in settings.get("layout_layers", []):
        if not layer.get("visible", True) or layer.get("role") == "watermark":
            continue
        y = float(layer.get("y_percent", 0))
        height = float(layer.get("height_percent", 0))
        if y < 35 and y + height <= 50:
            header_bottom = max(header_bottom or 0, y + height)
        if y >= 55:
            footer_top = min(footer_top if footer_top is not None else 100, y)
    if header_bottom is not None:
        top = max(top, math.ceil(header_bottom * height_mm / 100 + 4))
    if footer_top is not None:
        bottom = max(bottom, math.ceil((100 - footer_top) * height_mm / 100 + 4))
    return top, bottom


def render_brand_canvas(settings: dict, assets: dict[str, bytes], *, first_page: bool = False, dpi: int = REFERENCE_DPI) -> bytes:
    """Rasterize only the fixed branding plane; document body stays native/editable."""
    paper = settings.get("paper_size", "A4")
    width_mm, height_mm = (210.0, 297.0) if paper == "A4" else (215.9, 279.4)
    if not 72 <= dpi <= 200:
        raise ValueError("Resolução de timbrado inválida.")
    scale = dpi / 25.4
    size = (round(width_mm * scale), round(height_mm * scale))
    canvas = Image.new("RGBA", size, _rgb(settings.get("paper_color", "#FFFFFF")) + (255,))
    mode = settings.get("layout_mode", "structured")
    if mode not in {"structured", "reconstructed", "composed", "exact"}:
        raise ValueError("Modo de composição inválido.")
    if mode == "exact":
        asset_id = settings.get("background_asset_id")
        if not asset_id or str(asset_id) not in assets:
            raise ValueError("Fundo de papel timbrado não encontrado.")
        raw = assets[str(asset_id)]
        _, safe, _ = validate_reference("background.png" if raw.startswith(b"\x89PNG") else "background.jpg", raw, "background")
        with Image.open(io.BytesIO(safe)) as source:
            canvas.alpha_composite(ImageOps.fit(source.convert("RGBA"), size, method=Image.Resampling.LANCZOS))
    elif mode == "composed":
        _render_composed(canvas, settings, assets, first_page=first_page, dpi=dpi)
    elif mode == "reconstructed":
        draw = ImageDraw.Draw(canvas)
        utility = settings.get("utility_font_family", FONTS[0])
        font_path = _font_file(utility)
        heading_path = _font_file(settings.get("heading_font_family", FONTS[0]))
        if not font_path or not heading_path:
            raise RuntimeError("As fontes aprovadas não estão instaladas no servidor.")
        left = round(_number(settings, "margin_left_mm", 30, 10, 50) * scale)
        right = size[0] - round(_number(settings, "margin_right_mm", 20, 10, 50) * scale)
        alignment = settings.get("header_alignment", "left")
        if alignment not in _ALIGN:
            raise ValueError("Alinhamento inválido.")
        logo_id = settings.get("logo_asset_id")
        header_top = round(_number(settings, "header_top_mm", 10, 0, 60) * scale)
        text_top = header_top
        if logo_id:
            if str(logo_id) not in assets:
                raise ValueError("Logotipo não encontrado.")
            raw = assets[str(logo_id)]
            _, safe, _ = validate_reference("logo.png" if raw.startswith(b"\x89PNG") else "logo.jpg", raw, "logo")
            with Image.open(io.BytesIO(safe)) as picture:
                logo = picture.convert("RGBA")
                width = round(_number(settings, "logo_width_mm", 30, 10, 60) * scale)
                height = max(1, round(width * logo.height / logo.width))
                logo.thumbnail((width, height), Image.Resampling.LANCZOS)
                logo_top = round(_number(settings, "logo_top_mm", 8, 0, 60) * scale)
                x = left if alignment == "left" else right - logo.width if alignment == "right" else round((size[0] - logo.width) / 2)
                canvas.alpha_composite(logo, (x, logo_top))
                text_top = max(header_top, logo_top + logo.height + round(2 * scale))
        header_text = settings.get("first_header_text", "") if first_page and settings.get("different_first_page") else settings.get("header_text", "")
        header_size = _number(settings, "header_font_size_pt", 9, 6, 18)
        header_font = ImageFont.truetype(font_path, max(1, round(header_size * dpi / 72)))
        _tracked_text(draw, (left, text_top, right, round(height_mm * .28 * scale)), header_text, header_font,
                      _rgb(settings.get("primary_color", "#17324D")) + (255,), alignment,
                      _number(settings, "header_letter_spacing_pt", 0, 0, 5) * dpi / 72,
                      bool(settings.get("header_uppercase", False)))
        if settings.get("header_divider", True):
            width = (right - left) * _number(settings, "header_divider_width_percent", 100, 20, 100) / 100
            x = left if alignment == "left" else right - width if alignment == "right" else (left + right - width) / 2
            y = max(text_top + round((len(str(header_text).splitlines()) + 1) * header_size * dpi / 72 * 1.35), round(31 * scale))
            draw.line((round(x), y, round(x + width), y), fill=_rgb(settings.get("accent_color", "#8B6F47")) + (255,),
                      width=max(1, round(_number(settings, "header_divider_thickness_pt", .75, .25, 3) * dpi / 72)))
        watermark_id = settings.get("watermark_asset_id")
        watermark_text = _text(settings.get("watermark_text", ""), 160)
        opacity = _number(settings, "watermark_opacity", .12, 0, .3)
        if opacity and (watermark_id or watermark_text):
            if watermark_id:
                if str(watermark_id) not in assets:
                    raise ValueError("Imagem de marca d'água não encontrada.")
                raw = assets[str(watermark_id)]
                _, safe, _ = validate_reference("watermark.png" if raw.startswith(b"\x89PNG") else "watermark.jpg", raw, "watermark")
                with Image.open(io.BytesIO(safe)) as source:
                    mark = source.convert("RGBA")
            else:
                mark_font = ImageFont.truetype(heading_path, round(_number(settings, "watermark_font_size_pt", 100, 24, 180) * dpi / 72))
                bounds = mark_font.getbbox(watermark_text)
                mark = Image.new("RGBA", (max(1, bounds[2] - bounds[0] + 20), max(1, bounds[3] - bounds[1] + 20)))
                ImageDraw.Draw(mark).text((10 - bounds[0], 10 - bounds[1]), watermark_text, font=mark_font,
                                          fill=_rgb(settings.get("primary_color", "#17324D")) + (255,))
            target_width = round(_number(settings, "watermark_width_mm", 100, 20, 180) * scale)
            mark.thumbnail((target_width, round(height_mm * .65 * scale)), Image.Resampling.LANCZOS)
            if settings.get("watermark_position", "diagonal") == "diagonal":
                mark = mark.rotate(_number(settings, "watermark_rotation_deg", 35, -90, 90), expand=True, resample=Image.Resampling.BICUBIC)
            mark.putalpha(mark.getchannel("A").point(lambda alpha: round(alpha * opacity)))
            center_x = size[0] * _number(settings, "watermark_x_percent", 50, 0, 100) / 100
            center_y = size[1] * _number(settings, "watermark_y_percent", 50, 0, 100) / 100
            canvas.alpha_composite(mark, (round(center_x - mark.width / 2), round(center_y - mark.height / 2)))
        footer = _text(settings.get("footer_text", ""), 1000)
        footer_size = _number(settings, "footer_font_size_pt", 9, 6, 18)
        footer_font = ImageFont.truetype(font_path, max(1, round(footer_size * dpi / 72)))
        footer_alignment = settings.get("footer_alignment", "center")
        bottom = size[1] - round(_number(settings, "footer_bottom_mm", 8, 0, 60) * scale)
        line_count = max(1, len(footer.splitlines()))
        footer_top = bottom - round(line_count * footer_size * dpi / 72 * 1.4)
        if settings.get("footer_divider", True):
            width = (right - left) * _number(settings, "footer_divider_width_percent", 100, 20, 100) / 100
            x = left if footer_alignment == "left" else right - width if footer_alignment == "right" else (left + right - width) / 2
            draw.line((round(x), footer_top - round(5 * scale), round(x + width), footer_top - round(5 * scale)),
                      fill=_rgb(settings.get("accent_color", "#8B6F47")) + (255,),
                      width=max(1, round(_number(settings, "footer_divider_thickness_pt", .75, .25, 3) * dpi / 72)))
        _tracked_text(draw, (left, footer_top, right, bottom), footer, footer_font,
                      _rgb(settings.get("primary_color", "#17324D")) + (255,), footer_alignment,
                      _number(settings, "footer_letter_spacing_pt", 0, 0, 5) * dpi / 72,
                      bool(settings.get("footer_uppercase", False)))
    output = io.BytesIO()
    canvas.convert("RGB").save(output, format="PNG", optimize=True)
    return output.getvalue()


def _anchor_picture(paragraph, content: bytes, width, height) -> None:
    picture = paragraph.add_run().add_picture(io.BytesIO(content), width=width, height=height)
    anchor = picture._inline
    anchor.tag = qn("wp:anchor")
    for key, value in {"distT": "0", "distB": "0", "distL": "0", "distR": "0", "simplePos": "0", "relativeHeight": "0", "behindDoc": "1", "locked": "1", "layoutInCell": "1", "allowOverlap": "1"}.items():
        anchor.set(key, value)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.insert(0, simple)
    for index, axis in enumerate(("H", "V"), 1):
        position = OxmlElement("wp:position" + axis)
        position.set("relativeFrom", "page")
        offset = OxmlElement("wp:posOffset")
        offset.text = "0"
        position.append(offset)
        anchor.insert(index, position)
    anchor.insert(4, OxmlElement("wp:wrapNone"))


def _watermark(header, settings: dict, assets: dict[str, bytes]) -> None:
    image_id = settings.get("watermark_asset_id")
    text = _text(settings.get("watermark_text", ""), 160)
    opacity = _number(settings, "watermark_opacity", 0.08, 0, 0.3)
    if (not image_id and not text) or opacity == 0:
        return
    if image_id:
        if str(image_id) not in assets:
            raise ValueError("Imagem de marca d'água não encontrada.")
        raw = assets[str(image_id)]
        _, safe, _ = validate_reference("watermark.png" if raw.startswith(b"\x89PNG") else "watermark.jpg", raw, "watermark")
        with Image.open(io.BytesIO(safe)) as source:
            canvas = source.convert("RGBA")
    else:
        font_path = _font_file(settings.get("heading_font_family", FONTS[0]))
        if not font_path:
            raise RuntimeError("As fontes aprovadas não estão instaladas no servidor.")
        font = ImageFont.truetype(font_path, 100)
        bounds = font.getbbox(text)
        canvas = Image.new("RGBA", (max(bounds[2] - bounds[0] + 40, 1), max(bounds[3] - bounds[1] + 40, 1)))
        draw = ImageDraw.Draw(canvas)
        draw.text((20 - bounds[0], 20 - bounds[1]), text, font=font, fill=tuple(_color(settings, "primary_color", "#1F3A5F")) + (255,))
    canvas.thumbnail((1600, 1600))
    if settings.get("watermark_position", "diagonal") not in {"center", "diagonal"}:
        raise ValueError("Posição de marca d'água inválida.")
    if settings.get("watermark_position", "diagonal") == "diagonal":
        canvas = canvas.rotate(_number(settings, "watermark_rotation_deg", 35, -90, 90), expand=True,
                               resample=Image.Resampling.BICUBIC)
    canvas.putalpha(canvas.getchannel("A").point(lambda alpha: round(alpha * opacity)))
    buffer = io.BytesIO()
    canvas.save(buffer, format="PNG")
    paragraph = header.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    picture = paragraph.add_run().add_picture(io.BytesIO(buffer.getvalue()), width=Mm(_number(settings, "watermark_width_mm", 120, 20, 180)))
    inline = picture._inline
    inline.tag = qn("wp:anchor")
    for key, value in {"distT": "0", "distB": "0", "distL": "0", "distR": "0", "simplePos": "0", "relativeHeight": "0", "behindDoc": "1", "locked": "0", "layoutInCell": "1", "allowOverlap": "1"}.items():
        inline.set(key, value)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    inline.insert(0, simple)
    for index, axis in enumerate(("H", "V"), 1):
        position = OxmlElement("wp:position" + axis)
        position.set("relativeFrom", "page")
        align = OxmlElement("wp:align")
        align.text = "center"
        position.append(align)
        inline.insert(index, position)
    inline.insert(4, OxmlElement("wp:wrapNone"))


def _decorate_section(section, settings: dict, assets: dict[str, bytes]) -> None:
    mode = settings.get("layout_mode", "structured")
    scoped_layers = mode == "composed" and any(layer.get("page_scope", "all") != "all" for layer in settings.get("layout_layers", []))
    first_only = settings.get("background_scope", "all") == "first" and mode in {"reconstructed", "composed", "exact"}
    first = bool(settings.get("different_first_page", False) or first_only or scoped_layers)
    section.different_first_page_header_footer = first
    headers = [(section.header, settings.get("header_text", ""))]
    footers = [section.footer]
    if first:
        headers.append((section.first_page_header, settings.get("first_header_text", "")))
        footers.append(section.first_page_footer)
    for header, _text_value in headers:
        include_layers = not first_only or header is section.first_page_header
        background_settings = settings if include_layers else {**settings, "layout_mode": "structured"}
        paragraph = header.add_paragraph()
        paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(1)
        _anchor_picture(paragraph, render_brand_canvas(background_settings, assets, first_page=header is section.first_page_header), section.page_width, section.page_height)
    if mode in {"reconstructed", "composed", "exact"}:
        for footer in footers:
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if settings.get("page_numbers", True):
                paragraph.add_run("Página ")
                _field(paragraph, "PAGE")
                paragraph.add_run(" de ")
                _field(paragraph, "NUMPAGES")
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = settings.get("utility_font_family", settings.get("font_family", FONTS[0]))
        return
    usable_width = (section.page_width - section.left_margin - section.right_margin) / Mm(1)
    for header, text in headers:
        text = _text(text, 1000)
        lines = sum(max(1, math.ceil(len(line) / max(1, usable_width / 1.7))) for line in text.splitlines()) if text else 0
        text_height = lines * 4
        if text_height + 6 > section.top_margin / Mm(1):
            raise ValueError("Aumente a margem superior ou reduza o cabeçalho para evitar sobreposição.")
        paragraph = header.paragraphs[0]
        paragraph.alignment = _ALIGN[settings.get("header_alignment", "left")]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        logo_id = settings.get("logo_asset_id")
        if logo_id:
            if str(logo_id) not in assets:
                raise ValueError("Logotipo não encontrado.")
            raw = assets[str(logo_id)]
            _, logo, _ = validate_reference("logo.png" if raw.startswith(b"\x89PNG") else "logo.jpg", raw, "logo")
            with Image.open(io.BytesIO(logo)) as picture:
                max_height = section.top_margin / Mm(1) - text_height - 8
                if max_height < 4:
                    raise ValueError("Aumente a margem superior para acomodar o logotipo.")
                width = min(_number(settings, "logo_width_mm", 25, 5, 100), max_height * picture.width / picture.height)
            paragraph.add_run().add_picture(io.BytesIO(logo), width=Mm(width))
            if text:
                paragraph.add_run().add_break()
        run = paragraph.add_run(text)
        run.font.size, run.font.color.rgb = Pt(_number(settings, "header_font_size_pt", 9, 6, 18)), _color(settings, "primary_color", "#1F3A5F")
        if settings.get("header_divider", True):
            _paragraph_border(paragraph, "bottom", settings.get("accent_color", "#8B6F47"), _number(settings, "header_divider_thickness_pt", .75, .25, 3))
        _watermark(header, settings, assets)
    footer_text = _text(settings.get("footer_text", ""), 1000)
    lines = sum(max(1, math.ceil(len(line) / max(1, usable_width / 1.6))) for line in footer_text.splitlines()) if footer_text else 0
    if lines * 4 + (4 if settings.get("page_numbers", True) else 0) + 6 > section.bottom_margin / Mm(1):
        raise ValueError("Aumente a margem inferior ou reduza o rodapé para evitar sobreposição.")
    for footer in footers:
        paragraph = footer.paragraphs[0]
        paragraph.alignment = _ALIGN[settings.get("footer_alignment", "center")]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.add_run(footer_text)
        if settings.get("page_numbers", True):
            if footer_text:
                paragraph.add_run().add_break()
            paragraph.add_run("Página ")
            _field(paragraph, "PAGE")
            paragraph.add_run(" de ")
            _field(paragraph, "NUMPAGES")
        for run in paragraph.runs:
            run.font.size = Pt(_number(settings, "footer_font_size_pt", 9, 6, 18))
            run.font.name = settings.get("utility_font_family", settings.get("font_family", FONTS[0]))
        if settings.get("footer_divider", True):
            _paragraph_border(paragraph, "top", settings.get("accent_color", "#8B6F47"), _number(settings, "footer_divider_thickness_pt", .75, .25, 3))


def _body(document, content: str, content_format: str, settings: dict) -> None:
    lines = content.split("\n")
    index = 0
    while index < len(lines):
        line = lines[index].rstrip("\r")
        if content_format == "markdown":
            if "|" in line and index + 1 < len(lines) and re.fullmatch(r"\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*", lines[index + 1]):
                rows = [[cell.strip() for cell in line.strip().strip("|").split("|")]]
                index += 2
                while index < len(lines) and "|" in lines[index] and lines[index].strip():
                    rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                    index += 1
                if len(rows) > 200 or not 1 <= len(rows[0]) <= 8 or any(len(row) != len(rows[0]) for row in rows):
                    raise ValueError("Tabela deve ter até 8 colunas, 200 linhas e células consistentes.")
                table = document.add_table(rows=1, cols=len(rows[0]))
                table.style = "Table Grid"
                for col, cell in enumerate(rows[0]):
                    _inline(table.rows[0].cells[col].paragraphs[0], cell)
                    for run in table.rows[0].cells[col].paragraphs[0].runs:
                        run.bold = True
                repeat = OxmlElement("w:tblHeader")
                table.rows[0]._tr.get_or_add_trPr().append(repeat)
                for row in rows[1:]:
                    cells = table.add_row().cells
                    for col, cell in enumerate(row):
                        _inline(cells[col].paragraphs[0], cell)
                continue
            heading = re.match(r"^(#{1,3})\s+(.+)$", line)
            listing = re.match(r"^\s*(?:([-*])|(\d+)\.)\s+(.+)$", line)
            if heading:
                heading_text = heading.group(2).upper() if settings.get("heading_uppercase", False) else heading.group(2)
                _inline(document.add_heading("", level=len(heading.group(1))), heading_text)
            elif listing:
                # Literal labels preserve numbering/restarts and remain editable.
                paragraph = document.add_paragraph()
                _inline(paragraph, ("• " if listing.group(1) else listing.group(2) + ". ") + listing.group(3))
                paragraph.paragraph_format.left_indent = Mm(6)
                paragraph.paragraph_format.first_line_indent = Mm(-4)
            else:
                _inline(document.add_paragraph(), line)
        else:
            document.add_paragraph(line)
        index += 1


def _render_docx(title: str, content: str, settings: dict, assets: dict[str, bytes], content_format: str) -> bytes:
    _text(title, 300)
    _text(content, MAX_CONTENT_CHARS)
    if not content.strip() or len(content.splitlines()) > 5000 or content_format not in {"plain", "markdown"}:
        raise ValueError("Conteúdo vazio, extenso ou formato não permitido.")
    for key in ("font_family", "heading_font_family", "utility_font_family"):
        if settings.get(key, FONTS[0]) not in FONTS:
            raise ValueError("Fonte não disponível. Escolha uma das fontes da lista.")
    for key in ("header_alignment", "footer_alignment"):
        if settings.get(key, "left") not in _ALIGN:
            raise ValueError("Alinhamento inválido.")
    required_top, required_bottom = required_content_margins(settings)
    if required_top > 80 or required_bottom > 80:
        raise ValueError("O cabeçalho ou rodapé ocupa a área de texto. Reduza-o antes de exportar.")
    if float(settings.get("margin_top_mm", 30)) < required_top or float(settings.get("margin_bottom_mm", 25)) < required_bottom:
        raise ValueError("Ajuste as margens à área segura indicada antes de exportar.")
    document = Document()
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.title = title
    document.core_properties.created = document.core_properties.modified = datetime(2000, 1, 1, tzinfo=timezone.utc)
    normal = document.styles["Normal"]
    _font(normal, settings.get("font_family", FONTS[0]))
    normal.font.size = Pt(_number(settings, "body_size_pt", 12, 9, 18))
    normal.font.color.rgb = _color(settings, "text_color", "#202020")
    normal.paragraph_format.line_spacing = _number(settings, "line_spacing", 1.5, 1, 2.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        _font(style, settings.get("heading_font_family", FONTS[0]))
        style.font.color.rgb = _color(settings, "primary_color", "#1F3A5F")
        style.font.size = Pt(_number(settings, "heading_size_pt", 16, 12, 32))
        _letter_spacing(style, _number(settings, "heading_letter_spacing_pt", 0, 0, 3))
        style.paragraph_format.keep_with_next = True
    for name in ("Header", "Footer"):
        _font(document.styles[name], settings.get("utility_font_family", settings.get("font_family", FONTS[0])))
        document.styles[name].font.size = Pt(9)
    section = document.sections[0]
    paper = settings.get("paper_size", "A4")
    if paper not in {"A4", "LETTER"}:
        raise ValueError("Tamanho de papel inválido.")
    section.page_width, section.page_height = (Mm(210), Mm(297)) if paper == "A4" else (Mm(215.9), Mm(279.4))
    for key, attribute in (("top", "top_margin"), ("bottom", "bottom_margin"), ("left", "left_margin"), ("right", "right_margin")):
        setattr(section, attribute, Mm(_number(settings, f"margin_{key}_mm", 25, 10, 80 if key in {"top", "bottom"} else 50)))
    section.header_distance = section.footer_distance = Mm(3)
    _decorate_section(section, settings, assets)
    if title.strip() and settings.get("show_document_title", True):
        document.add_heading(title.upper() if settings.get("heading_uppercase", False) else title, 0)
    _body(document, content, content_format, settings)
    buffer = io.BytesIO()
    document.save(buffer)
    if buffer.tell() > MAX_OUTPUT_BYTES:
        raise ValueError("Documento gerado excede o limite de tamanho.")
    return buffer.getvalue()


def _convert_pdf(docx: bytes) -> bytes:
    binary = _office_binary()
    if not binary or not all(_font_file(font) for font in FONTS):
        raise RuntimeError("Exportação PDF indisponível: LibreOffice e as fontes aprovadas precisam estar instalados.")
    with tempfile.TemporaryDirectory(prefix="lexflow-brand-") as directory:
        work = Path(directory)
        source = work / "document.docx"
        source.write_bytes(docx)
        command = [binary, "-env:UserInstallation=" + (work / "profile").as_uri(), "--headless", "--nologo", "--nodefault", "--nolockcheck", "--norestore", "--convert-to", 'pdf:writer_pdf_Export:{"UseTaggedPDF":{"type":"boolean","value":"true"},"ExportBookmarks":{"type":"boolean","value":"true"}}', "--outdir", str(work), str(source)]
        # prlimit applies limits before exec without unsafe preexec_fn in an API
        # thread. The runtime image provides it; Windows local mode has timeout.
        if os.name == "posix":
            limiter = shutil.which("prlimit")
            if not limiter:
                raise RuntimeError("O servidor precisa de prlimit para limitar a conversão.")
            command = [limiter, "--cpu=40", "--as=1073741824", f"--fsize={MAX_OUTPUT_BYTES}", "--nofile=256", "--"] + command
        environment = {key: value for key, value in os.environ.items() if key in {"PATH", "SYSTEMROOT", "WINDIR", "LANG", "LC_ALL"}}
        environment.update({"HOME": directory, "TMPDIR": directory, "SAL_USE_VCLPLUGIN": "svp", "SAL_DISABLE_OPENCL": "1"})
        try:
            process = subprocess.Popen(command, cwd=directory, env=environment, stdin=subprocess.DEVNULL, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, start_new_session=os.name == "posix", creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0)
        except OSError as exc:
            raise RuntimeError("Não foi possível iniciar a conversão PDF no servidor.") from exc
        try:
            process.wait(timeout=RENDER_TIMEOUT)
        except subprocess.TimeoutExpired as exc:
            if os.name == "posix":
                try:
                    os.killpg(process.pid, signal.SIGKILL)
                except ProcessLookupError:
                    pass
            else:
                process.kill()
            process.wait(timeout=5)
            raise RuntimeError("A exportação excedeu o tempo limite; reduza o documento e tente novamente.") from exc
        output = work / "document.pdf"
        if process.returncode != 0 or not output.is_file() or not 0 < output.stat().st_size <= MAX_OUTPUT_BYTES:
            raise RuntimeError("Não foi possível gerar o PDF com segurança.")
        pdf = output.read_bytes()
        try:
            reader = PdfReader(io.BytesIO(pdf), strict=True)
            if reader.is_encrypted or not 1 <= len(reader.pages) <= MAX_PAGES:
                raise ValueError("PDF inválido ou acima do limite de páginas.")
        except Exception as exc:
            raise RuntimeError("O PDF gerado não passou na validação.") from exc
        return pdf


def render_documents(title: str, content: str, settings: dict, assets: dict[str, bytes], content_format: str = "plain") -> tuple[bytes, bytes]:
    """Return editable DOCX and searchable PDF; never return a fake PDF fallback."""
    # ponytail: one conversion per process; use a dedicated queue when measured
    # demand exceeds the two API workers configured for the initial VPS.
    if not _RENDER_SLOTS.acquire(timeout=1):
        raise RuntimeError("Outra exportação está em andamento. Tente novamente em instantes.")
    try:
        docx = _render_docx(title, content, settings, assets, content_format)
        return docx, _convert_pdf(docx)
    finally:
        _RENDER_SLOTS.release()


def render_docx_only(title: str, content: str, settings: dict, assets: dict[str, bytes], content_format: str = "plain") -> bytes:
    """Return editable DOCX directly without invoking LibreOffice PDF conversion."""
    return _render_docx(title, content, settings, assets, content_format)
